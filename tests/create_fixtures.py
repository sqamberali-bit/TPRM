"""Generate test fixture files (Word report).

Run once:  python tests/create_fixtures.py
"""
from pathlib import Path

from docx import Document

FIXTURES = Path(__file__).parent / "fixtures"
FIXTURES.mkdir(exist_ok=True)


def create_sample_report() -> None:
    doc = Document()
    doc.add_heading("Acme Corp Security Assessment", level=1)
    doc.add_paragraph("Date: 2025-06-15")

    # --- Consolidated findings table ---
    t1 = doc.add_table(rows=4, cols=4)
    t1.style = "Table Grid"
    headers = ["Ref", "Finding", "Rating", "Basis"]
    for i, h in enumerate(headers):
        t1.rows[0].cells[i].text = h

    data = [
        ("F1", "No MFA on admin accounts", "High", "Admin portal lacks MFA"),
        ("F2", "Unencrypted data at rest", "Medium", "Database not encrypted"),
        ("F3", "Outdated TLS configuration", "Low", "TLS 1.0 still enabled"),
    ]
    for row_idx, (ref, finding, rating, basis) in enumerate(data, 1):
        t1.rows[row_idx].cells[0].text = ref
        t1.rows[row_idx].cells[1].text = finding
        t1.rows[row_idx].cells[2].text = rating
        t1.rows[row_idx].cells[3].text = basis

    doc.add_paragraph("")

    # --- Residual risks table ---
    t2 = doc.add_table(rows=4, cols=5)
    t2.style = "Table Grid"
    headers2 = [
        "Ref",
        "Residual Risk",
        "Compensating Controls",
        "Rating",
        "Duration / Review Trigger",
    ]
    for i, h in enumerate(headers2):
        t2.rows[0].cells[i].text = h

    data2 = [
        ("F1", "Admin access without MFA", "IP whitelist in place", "Medium",
         "Accepted until December 2025"),
        ("F2", "Data at rest unencrypted", "Network segmentation", "Low",
         "Reverts to High if clause F9 not executed"),
        ("F3", "TLS 1.0 enabled", "WAF in front", "Low", "Closed"),
    ]
    for row_idx, vals in enumerate(data2, 1):
        for col_idx, val in enumerate(vals):
            t2.rows[row_idx].cells[col_idx].text = val

    doc.add_paragraph("")

    # --- Outstanding items table ---
    t3 = doc.add_table(rows=3, cols=4)
    t3.style = "Table Grid"
    headers3 = ["Item", "Resolution", "Status", "Ref"]
    for i, h in enumerate(headers3):
        t3.rows[0].cells[i].text = h

    data3 = [
        ("Enable MFA on admin portal", "Vendor committed to Q3 rollout",
         "In progress", "F1"),
        ("Encrypt database at rest", "Clause added to contract",
         "Accepted", "F2"),
    ]
    for row_idx, vals in enumerate(data3, 1):
        for col_idx, val in enumerate(vals):
            t3.rows[row_idx].cells[col_idx].text = val

    path = FIXTURES / "sample_report.docx"
    doc.save(str(path))
    print(f"Created {path}")


if __name__ == "__main__":
    create_sample_report()
