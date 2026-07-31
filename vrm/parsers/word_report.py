"""Parse vendor security assessment Word reports.

Extracts three table types:
1. Consolidated findings (Ref, Finding, Rating, Basis)
2. Residual risks (Ref, Residual risk, Compensating controls, Rating, Duration/review trigger)
3. Historical outstanding items (Item, Resolution, Status, Ref)

Matches by column header with fuzzy tolerance for minor wording changes.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from docx import Document
from docx.table import Table

from vrm.utils.logging import get_logger

log = get_logger(__name__)


@dataclass
class Finding:
    ref: str
    title: str
    rating: str
    basis: str = ""


@dataclass
class ResidualRisk:
    ref: str
    title: str
    compensating_controls: str = ""
    rating: str = ""
    duration_or_trigger: str = ""


@dataclass
class OutstandingItem:
    item: str
    resolution: str = ""
    status: str = ""
    ref: str = ""


@dataclass
class ParsedReport:
    source_file: str
    vendor_name: Optional[str] = None
    report_date: Optional[str] = None
    findings: list[Finding] = field(default_factory=list)
    residual_risks: list[ResidualRisk] = field(default_factory=list)
    outstanding_items: list[OutstandingItem] = field(default_factory=list)
    unmatched_tables: int = 0


_FINDINGS_HEADERS = {"ref", "finding", "rating", "basis"}
_FINDINGS_ALIASES = {
    "reference": "ref",
    "finding reference": "ref",
    "finding title": "finding",
    "risk rating": "rating",
    "inherent rating": "rating",
    "basis / rationale": "basis",
    "rationale": "basis",
}

_RESIDUAL_HEADERS = {"ref", "residual risk", "compensating controls", "rating"}
_RESIDUAL_ALIASES = {
    "reference": "ref",
    "residual rating": "rating",
    "residual risk rating": "rating",
    "compensating control": "compensating controls",
    "controls": "compensating controls",
    "duration": "duration/review trigger",
    "duration / review trigger": "duration/review trigger",
    "review trigger": "duration/review trigger",
    "duration/trigger": "duration/review trigger",
}

_OUTSTANDING_HEADERS = {"item", "resolution", "status"}
_OUTSTANDING_ALIASES = {
    "outstanding item": "item",
    "finding": "item",
    "description": "item",
    "resolution / status": "resolution",
    "current status": "status",
    "reference": "ref",
}


def _normalise_header(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip().lower())


def _match_headers(
    row_cells: list[str],
    required: set[str],
    aliases: dict[str, str],
) -> Optional[dict[int, str]]:
    """Try to map column indices to canonical field names.

    Returns a mapping {col_index: field_name} if enough required headers
    match, otherwise None.
    """
    mapping: dict[int, str] = {}
    for i, cell in enumerate(row_cells):
        norm = _normalise_header(cell)
        if norm in required or norm in {"duration/review trigger", "ref"}:
            mapping[i] = norm
        elif norm in aliases:
            mapping[i] = aliases[norm]

    matched_fields = set(mapping.values())
    core_required = required - {"duration/review trigger"}
    if core_required - matched_fields:
        return None
    return mapping


def _table_rows(table: Table) -> list[list[str]]:
    rows = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])
    return rows


def _classify_and_parse_table(
    table: Table,
) -> Optional[list[Finding] | list[ResidualRisk] | list[OutstandingItem]]:
    rows = _table_rows(table)
    if len(rows) < 2:
        return None

    header_row = rows[0]

    residual_map = _match_headers(header_row, _RESIDUAL_HEADERS, _RESIDUAL_ALIASES)
    if residual_map:
        return _parse_residual_table(rows[1:], residual_map)

    findings_map = _match_headers(header_row, _FINDINGS_HEADERS, _FINDINGS_ALIASES)
    if findings_map:
        return _parse_findings_table(rows[1:], findings_map)

    outstanding_map = _match_headers(
        header_row, _OUTSTANDING_HEADERS, _OUTSTANDING_ALIASES
    )
    if outstanding_map:
        return _parse_outstanding_table(rows[1:], outstanding_map)

    return None


def _cell(row: list[str], mapping: dict[int, str], field_name: str) -> str:
    for idx, name in mapping.items():
        if name == field_name:
            return row[idx] if idx < len(row) else ""
    return ""


def _parse_findings_table(
    rows: list[list[str]], mapping: dict[int, str]
) -> list[Finding]:
    results = []
    for row in rows:
        ref = _cell(row, mapping, "ref")
        if not ref:
            continue
        results.append(
            Finding(
                ref=ref,
                title=_cell(row, mapping, "finding"),
                rating=_cell(row, mapping, "rating"),
                basis=_cell(row, mapping, "basis"),
            )
        )
    return results


def _parse_residual_table(
    rows: list[list[str]], mapping: dict[int, str]
) -> list[ResidualRisk]:
    results = []
    for row in rows:
        ref = _cell(row, mapping, "ref")
        if not ref:
            continue
        results.append(
            ResidualRisk(
                ref=ref,
                title=_cell(row, mapping, "residual risk"),
                compensating_controls=_cell(row, mapping, "compensating controls"),
                rating=_cell(row, mapping, "rating"),
                duration_or_trigger=_cell(row, mapping, "duration/review trigger"),
            )
        )
    return results


def _parse_outstanding_table(
    rows: list[list[str]], mapping: dict[int, str]
) -> list[OutstandingItem]:
    results = []
    for row in rows:
        item = _cell(row, mapping, "item")
        if not item:
            continue
        results.append(
            OutstandingItem(
                item=item,
                resolution=_cell(row, mapping, "resolution"),
                status=_cell(row, mapping, "status"),
                ref=_cell(row, mapping, "ref"),
            )
        )
    return results


def parse_report(path: Path) -> ParsedReport:
    log.info("parsing_word_report", path=str(path))
    doc = Document(str(path))
    result = ParsedReport(source_file=path.name)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not result.vendor_name and para.style.name.startswith("Heading"):
            result.vendor_name = text
            break

    for table in doc.tables:
        parsed = _classify_and_parse_table(table)
        if parsed is None:
            result.unmatched_tables += 1
            log.debug("unmatched_table", file=path.name)
            continue

        if parsed and isinstance(parsed[0], Finding):
            result.findings.extend(parsed)
            log.info("parsed_findings", count=len(parsed), file=path.name)
        elif parsed and isinstance(parsed[0], ResidualRisk):
            result.residual_risks.extend(parsed)
            log.info("parsed_residual_risks", count=len(parsed), file=path.name)
        elif parsed and isinstance(parsed[0], OutstandingItem):
            result.outstanding_items.extend(parsed)
            log.info("parsed_outstanding_items", count=len(parsed), file=path.name)

    if result.unmatched_tables:
        log.warning(
            "unmatched_tables_in_report",
            count=result.unmatched_tables,
            file=path.name,
        )

    return result
